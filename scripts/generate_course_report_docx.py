#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert 数据库课程设计报告.md to professional Word (.docx)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "08-课设交付" / "数据库课程设计报告.md"
OUT = ROOT / "docs" / "08-课设交付" / "数据库课程设计报告.docx"


def set_doc_font(doc: Document, name: str = "宋体", size: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def parse_md_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-:\s|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            table.rows[ri].cells[ci].text = text


def is_heading(line: str) -> tuple[int, str] | None:
    m = re.match(r"^(#{1,4})\s+(.+)$", line.strip())
    if not m:
        return None
    level = len(m.group(1))
    return level, m.group(2).strip()


def convert(md_text: str) -> Document:
    doc = Document()
    set_doc_font(doc)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            add_table(doc, parse_md_table(table_buf))
            table_buf = []

    def flush_code() -> None:
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            code_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        if stripped in ("---", ""):
            i += 1
            continue

        h = is_heading(stripped)
        if h:
            level, text = h
            if level == 1 and text == "数据库课程设计报告":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(22)
            elif level == 2 and text.startswith("校园自习室"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(16)
            else:
                doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        if stripped.startswith("**关键词**"):
            doc.add_paragraph(stripped.replace("**", ""))
            i += 1
            continue

        text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = text.replace("组长验收", "课程验收").replace("组长", "验收方")
        doc.add_paragraph(text)
        i += 1

    flush_table()
    flush_code()
    return doc


def main() -> None:
    if not MD.exists():
        raise FileNotFoundError(MD)
    text = MD.read_text(encoding="utf-8")
    doc = convert(text)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"written: {OUT}")


if __name__ == "__main__":
    main()

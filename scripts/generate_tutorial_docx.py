#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 docs/09-理解与讲解 下的 Word 文档。

用法:
  python scripts/generate_tutorial_docx.py
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "09-理解与讲解"

DOCS = [
    (
        OUT_DIR / "01-项目理解指南.md",
        OUT_DIR / "01-项目理解指南.docx",
        "校园自习室预约管理系统",
        "01 · 项目理解指南（零基础必读）",
    ),
    (
        OUT_DIR / "02-答辩讲解手册.md",
        OUT_DIR / "02-答辩讲解手册.docx",
        "校园自习室预约管理系统",
        "02 · 答辩讲解手册（演示+指代码）",
    ),
]

BODY_FONT = "宋体"
CODE_FONT = "Consolas"
BODY_SIZE = Pt(12)
TITLE_SIZE = Pt(22)
SUBTITLE_SIZE = Pt(16)
TAG_COLORS = {
    "【讲】": RGBColor(0x1A, 0x56, 0xDB),
    "【指】": RGBColor(0x0D, 0x7A, 0x3B),
    "【演】": RGBColor(0xB4, 0x5A, 0x00),
    "【实例】": RGBColor(0x5B, 0x2C, 0x8F),
    "【定义】": RGBColor(0x0E, 0x66, 0x55),
    "【在本项目】": RGBColor(0x33, 0x33, 0x33),
    "【要看代码】": RGBColor(0xC0, 0x39, 0x2B),
}


def set_run_font(run, name: str, size: Pt | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element.get_or_add_rPr()
    rfonts = r.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)


def set_doc_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.74)


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def parse_md_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[-:\s|]+\|$", line):
            continue
        rows.append([c.strip() for c in line.strip("|").split("|")])
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            cell.text = clean_inline(row[ci] if ci < len(row) else "")
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    set_run_font(run, BODY_FONT, Pt(10.5))
            if ri == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def add_paragraph(doc: Document, text: str, *, indent: bool = True, bold: bool = False) -> None:
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    for tag, color in TAG_COLORS.items():
        if text.startswith(tag):
            r1 = p.add_run(tag + " ")
            set_run_font(r1, BODY_FONT, BODY_SIZE, True, color)
            r2 = p.add_run(clean_inline(text[len(tag):].strip()))
            set_run_font(r2, BODY_FONT, BODY_SIZE, bold)
            return
    run = p.add_run(clean_inline(text))
    set_run_font(run, BODY_FONT, BODY_SIZE, bold)


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, CODE_FONT, Pt(9.5))


def add_cover(doc: Document, subtitle: str, doc_title: str) -> None:
    for line, size in [("华南农业大学 · 数据库课程设计", 14), (subtitle, 18)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        set_run_font(p.add_run(line), "黑体", Pt(size), True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    p3.paragraph_format.first_line_indent = Cm(0)
    set_run_font(p3.add_run(doc_title), "黑体", SUBTITLE_SIZE, True)

    for label, value in [
        ("文档目录", "docs/09-理解与讲解"),
        ("项目路径", r"D:\SchoolWorkPlace\Database\CSRRMS"),
        ("导航入口", "00-文档导航.md / 架构与文档导航.md"),
    ]:
        px = doc.add_paragraph()
        px.paragraph_format.first_line_indent = Cm(0)
        px.paragraph_format.left_indent = Cm(2.5)
        set_run_font(px.add_run(f"{label}："), BODY_FONT, BODY_SIZE, True)
        set_run_font(px.add_run(value), BODY_FONT, BODY_SIZE)

    doc.add_page_break()


def heading_level(md_level: int, text: str) -> int:
    if re.match(r"^第\s*\d+\s*章", text) or re.match(r"^第\s*[一二三四五六七八九十]+", text):
        return 1
    if text.startswith("第") and "节" in text[:6]:
        return 1
    if re.match(r"^\d+\.\d+", text):
        return 2
    if md_level <= 2:
        return 1
    if md_level == 3:
        return 2
    return 3


def convert_md(md_text: str, subtitle: str, doc_title: str) -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc, subtitle, doc_title)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []
    skip_h1 = True

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            add_table(doc, parse_md_table(table_buf))
            table_buf = []

    def flush_code() -> None:
        nonlocal code_buf
        if code_buf:
            add_code_block(doc, code_buf)
            code_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if stripped.startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        if stripped in ("---", ""):
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if m:
            level, text = len(m.group(1)), m.group(2).strip()
            if skip_h1 and level == 1:
                skip_h1 = False
                i += 1
                continue
            hl = doc.add_heading(clean_inline(text), level=heading_level(level, text))
            hl.paragraph_format.first_line_indent = Cm(0)
            for run in hl.runs:
                set_run_font(run, "黑体")
            i += 1
            continue

        if stripped.startswith("> "):
            add_paragraph(doc, stripped[2:])
            i += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            set_run_font(p.add_run(clean_inline(stripped[2:])), BODY_FONT, BODY_SIZE)
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    flush_table()
    flush_code()
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    archive_dir = ROOT.parent / "课设资料" / "06-理解与讲解"
    sync_names = {
        "01-项目理解指南.docx": "01-项目理解指南.docx",
        "02-答辩讲解手册.docx": "02-答辩讲解手册.docx",
    }
    for md_path, out_path, subtitle, doc_title in DOCS:
        if not md_path.exists():
            raise FileNotFoundError(md_path)
        doc = convert_md(md_path.read_text(encoding="utf-8"), subtitle, doc_title)
        doc.save(str(out_path))
        print(f"written: {out_path} ({out_path.stat().st_size} bytes)")
        if archive_dir.parent.exists():
            archive_dir.mkdir(parents=True, exist_ok=True)
            key = out_path.name
            if key in sync_names:
                dest = archive_dir / sync_names[key]
                shutil.copy2(out_path, dest)
                print(f"archived: {dest}")


if __name__ == "__main__":
    main()
